"""
RAG Provider Abstraction Layer
Supports 3 modes: google, local, ultra
Controlled by RAG_MODE environment variable
"""

import os
import re
import json
import time
import hashlib
import asyncio
import sqlite3
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Tuple, Generator
from datetime import datetime

# ============================================================================
# MODE CONFIGURATION
# ============================================================================
RAG_MODE = os.environ.get("RAG_MODE", "google")  # google | local | ultra

# ============================================================================
# BASE PROVIDER INTERFACE
# ============================================================================
class BaseRetriever(ABC):
    """Base class for document retrieval."""

    @abstractmethod
    async def retrieve(self, query: str) -> Tuple[List[str], List[Dict[str, Any]]]:
        """Retrieve relevant documents. Returns (contexts, sources)."""
        pass


class BaseReranker(ABC):
    """Base class for document reranking."""

    @abstractmethod
    def rerank(self, query: str, documents: List[Dict[str, Any]], top_n: int = 5) -> List[Dict[str, Any]]:
        """Rerank documents by relevance."""
        pass


class BaseGenerator(ABC):
    """Base class for LLM response generation."""

    @abstractmethod
    def generate(self, prompt: str) -> str:
        """Generate a response."""
        pass

    @abstractmethod
    def generate_stream(self, prompt: str) -> Generator[str, None, None]:
        """Generate a streaming response."""
        pass


# ============================================================================
# GOOGLE MODE PROVIDERS (uses existing app.py functions)
# ============================================================================
# Google providers are thin wrappers — the actual logic stays in app.py
# to avoid duplicating the Google Cloud SDK setup


# ============================================================================
# LOCAL MODE PROVIDERS
# ============================================================================
_chroma_collection = None
_embedding_model = None
_local_docs_ingested = False


def get_embedding_model():
    """Lazy-load sentence-transformers model."""
    global _embedding_model
    if _embedding_model is None:
        from sentence_transformers import SentenceTransformer
        model_name = os.environ.get(
            "LOCAL_EMBEDDING_MODEL",
            "mesolitica/mistral-embedding-191m-8k-contrastive"
        )
        print(f"[LOCAL] Loading embedding model: {model_name}")
        _embedding_model = SentenceTransformer(model_name)
        print(f"[LOCAL] Embedding model loaded successfully")
    return _embedding_model


def get_chroma_collection():
    """Get or create ChromaDB collection."""
    global _chroma_collection
    if _chroma_collection is None:
        import chromadb

        persist_dir = os.environ.get(
            "CHROMA_PERSIST_DIR",
            "/opt/jkst-master-ai/chroma_db"
        )
        print(f"[LOCAL] Initializing ChromaDB at: {persist_dir}")
        client = chromadb.PersistentClient(path=persist_dir)
        _chroma_collection = client.get_or_create_collection(
            name="jkst_knowledge",
            metadata={"hnsw:space": "cosine"}
        )
        print(f"[LOCAL] ChromaDB collection ready: {_chroma_collection.count()} documents")
    return _chroma_collection


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        import fitz  # pymupdf
        doc = fitz.open(filepath)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"[LOCAL] PDF extraction error for {filepath}: {e}")
        return ""


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"[LOCAL] DOCX extraction error for {filepath}: {e}")
        return ""


def extract_text_from_file(filepath: str) -> str:
    """Extract text from any supported file type."""
    ext = filepath.rsplit('.', 1)[-1].lower() if '.' in filepath else ''

    if ext == 'pdf':
        return extract_text_from_pdf(filepath)
    elif ext == 'docx':
        return extract_text_from_docx(filepath)
    elif ext in ('md', 'txt'):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception as e:
            print(f"[LOCAL] Text read error for {filepath}: {e}")
            return ""
    else:
        return ""


def ingest_knowledge_to_chroma(knowledge_path: str = "/opt/jkst-master-ai/knowledge"):
    """Ingest local knowledge files + documents into ChromaDB with chunking."""
    global _local_docs_ingested
    if _local_docs_ingested:
        return

    import glob as globlib

    collection = get_chroma_collection()

    # Skip if already has documents
    if collection.count() > 0:
        print(f"[LOCAL] ChromaDB already has {collection.count()} chunks, skipping ingestion")
        _local_docs_ingested = True
        return

    model = get_embedding_model()

    # Gather all supported files from knowledge folder
    all_files = []
    for ext in ('*.md', '*.txt', '*.pdf', '*.docx'):
        all_files += globlib.glob(os.path.join(knowledge_path, ext))

    # Also ingest from local documents folder if it exists
    local_docs_path = os.environ.get("LOCAL_DOCUMENTS_PATH", "/opt/jkst-master-ai/documents")
    if os.path.exists(local_docs_path):
        for ext in ('*.md', '*.txt', '*.pdf', '*.docx'):
            all_files += globlib.glob(os.path.join(local_docs_path, ext))

    print(f"[LOCAL] Found {len(all_files)} files to ingest")

    chunk_size = 500  # characters per chunk
    chunk_overlap = 100

    all_ids = []
    all_texts = []
    all_metadatas = []
    all_embeddings = []
    skipped = 0

    for filepath in all_files:
        try:
            content = extract_text_from_file(filepath)
            if not content or len(content) < 50:
                skipped += 1
                continue

            filename = os.path.basename(filepath)
            file_ext = filepath.rsplit('.', 1)[-1].lower()

            # Determine document type label
            if file_ext == 'pdf':
                doc_type = "Dokumen PDF"
            elif file_ext == 'docx':
                doc_type = "Dokumen DOCX"
            else:
                doc_type = "Fail Pengetahuan Tempatan"

            # Split into chunks
            chunks = []
            for i in range(0, len(content), chunk_size - chunk_overlap):
                chunk = content[i:i + chunk_size]
                if len(chunk.strip()) > 50:  # Skip very small chunks
                    chunks.append(chunk)

            # Prepare for embedding
            for idx, chunk in enumerate(chunks):
                doc_id = hashlib.md5(f"{filename}:{idx}".encode()).hexdigest()
                all_ids.append(doc_id)
                all_texts.append(chunk)
                all_metadatas.append({
                    "filename": filename,
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                    "source": filepath,
                    "type": doc_type
                })

            print(f"[LOCAL] Chunked {filename}: {len(chunks)} chunks ({len(content)} chars extracted)")

        except Exception as e:
            print(f"[LOCAL] Error processing {filepath}: {e}")
            skipped += 1

    if skipped:
        print(f"[LOCAL] Skipped {skipped} files (empty or unreadable)")

    if all_texts:
        total = len(all_texts)
        embed_batch_size = 256  # Embed in batches to avoid OOM and show progress
        store_batch_size = 100  # ChromaDB storage batch

        print(f"[LOCAL] Embedding {total} chunks in batches of {embed_batch_size}...")

        all_embeddings = []
        for batch_start in range(0, total, embed_batch_size):
            batch_end = min(batch_start + embed_batch_size, total)
            batch_texts = all_texts[batch_start:batch_end]
            batch_embs = model.encode(batch_texts, show_progress_bar=False).tolist()
            all_embeddings.extend(batch_embs)
            progress = len(all_embeddings)
            print(f"[LOCAL] Embedded {progress}/{total} chunks ({progress*100//total}%)")

        # Add to ChromaDB in batches
        for i in range(0, len(all_ids), store_batch_size):
            end = min(i + store_batch_size, len(all_ids))
            collection.add(
                ids=all_ids[i:end],
                documents=all_texts[i:end],
                metadatas=all_metadatas[i:end],
                embeddings=all_embeddings[i:end]
            )

        print(f"[LOCAL] Ingested {total} chunks into ChromaDB")

    _local_docs_ingested = True


class LocalRetriever(BaseRetriever):
    """Retrieves from local ChromaDB using sentence-transformer embeddings."""

    def __init__(self):
        self.top_k = int(os.environ.get("LOCAL_RETRIEVAL_TOP_K", "10"))

    async def retrieve(self, query: str) -> Tuple[List[str], List[Dict[str, Any]]]:
        # Ensure knowledge is ingested
        await asyncio.to_thread(ingest_knowledge_to_chroma)

        collection = get_chroma_collection()
        model = get_embedding_model()

        # Embed query
        query_embedding = await asyncio.to_thread(model.encode, query)

        # Search ChromaDB
        results = collection.query(
            query_embeddings=[query_embedding.tolist()],
            n_results=self.top_k,
            include=["documents", "metadatas", "distances"]
        )

        contexts = []
        sources = []

        if results and results['documents']:
            for i, (doc, meta, dist) in enumerate(zip(
                results['documents'][0],
                results['metadatas'][0],
                results['distances'][0]
            )):
                score = 1.0 - dist  # cosine distance to similarity
                filename = meta.get('filename', 'Unknown')

                contexts.append(
                    f"[SUMBER - DOKUMEN TEMPATAN: {filename}]\n{doc}"
                )
                sources.append({
                    "type": "Dokumen Tempatan (ChromaDB)",
                    "filename": filename,
                    "page_content": doc[:500],
                    "score": round(score, 4),
                    "source_uri": f"local://{filename}",
                    "chunk_index": meta.get('chunk_index', 0),
                    "priority": "PRIMARY"
                })

        return contexts, sources


class LocalReranker(BaseReranker):
    """Simple TF-IDF based reranker for local mode (no external API)."""

    def rerank(self, query: str, documents: List[Dict[str, Any]], top_n: int = 5) -> List[Dict[str, Any]]:
        if not documents or len(documents) <= top_n:
            return documents

        query_words = set(query.lower().split())

        scored = []
        for doc in documents:
            text = doc.get("page_content", "").lower()
            # Simple keyword overlap scoring
            score = sum(1 for w in query_words if w in text and len(w) > 2)
            # Boost by original vector score
            score += doc.get("score", 0) * 5
            scored.append((score, doc))

        scored.sort(key=lambda x: x[0], reverse=True)
        result = [doc for _, doc in scored[:top_n]]
        print(f"[LOCAL] Simple reranking: {len(documents)} → {len(result)} docs")
        return result


class GeminiGenerator(BaseGenerator):
    """Uses Gemini for generation (shared by all modes for now)."""

    def __init__(self):
        self._client = None
        self.model = os.environ.get("GCP_GEMINI_MODEL", "gemini-2.0-flash-001")

    def _get_client(self):
        if self._client is None:
            from google import genai
            from google.oauth2 import service_account
            import google.auth.transport.requests

            creds_path = os.environ.get(
                "GOOGLE_APPLICATION_CREDENTIALS",
                "/home/adilhidayat/jkst-credentials.json"
            )
            project = os.environ.get("GCP_PROJECT_ID", "jab-kehakiman-syariah-tgg")
            location = os.environ.get("GCP_GEMINI_LOCATION", "us-central1")

            scopes = ['https://www.googleapis.com/auth/cloud-platform']
            credentials = service_account.Credentials.from_service_account_file(
                creds_path, scopes=scopes
            )
            if not credentials.valid:
                credentials.refresh(google.auth.transport.requests.Request())

            self._client = genai.Client(
                vertexai=True,
                project=project,
                location=location,
                credentials=credentials
            )
        return self._client

    def generate(self, prompt: str) -> str:
        client = self._get_client()
        response = client.models.generate_content(
            model=self.model,
            contents=prompt,
            config={"temperature": 0.2, "max_output_tokens": 2000}
        )
        return response.text

    def generate_stream(self, prompt: str) -> Generator[str, None, None]:
        client = self._get_client()
        stream = client.models.generate_content_stream(
            model=self.model,
            contents=prompt,
            config={"temperature": 0.2, "max_output_tokens": 2000}
        )
        for chunk in stream:
            if chunk.text:
                yield chunk.text


class OpenAICompatibleGenerator(BaseGenerator):
    """
    Uses any OpenAI-compatible API (YTL Ilmu, Ollama, etc.).

    YTL Ilmu quirk: returns empty content without tools.
    Workaround: always send a dummy tool so the model activates,
    then handle tool calls with a follow-up message.
    """

    # Dummy tool that tricks YTL into generating content
    DUMMY_TOOL = {
        "type": "function",
        "function": {
            "name": "respond_to_user",
            "description": "Respond directly to the user with a helpful answer. Use this tool to provide your answer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "answer": {"type": "string", "description": "Your complete answer to the user's question"}
                },
                "required": ["answer"],
            },
        },
    }

    def __init__(self):
        self.base_url = os.environ.get("LOCAL_LLM_BASE_URL", "https://api.ytlailabs.tech/v1")
        self.api_key = os.environ.get("LOCAL_LLM_API_KEY", "")
        self.model = os.environ.get("LOCAL_LLM_MODEL", "ilmu-text-free-v2")
        self.use_tool_workaround = os.environ.get("LOCAL_LLM_TOOL_WORKAROUND", "true").lower() == "true"
        self._client = None

    def _get_client(self):
        if self._client is None:
            from openai import OpenAI
            self._client = OpenAI(
                base_url=self.base_url,
                api_key=self.api_key
            )
        return self._client

    def _call_with_tool_workaround(self, messages: list, max_tokens: int = 2000, temperature: float = 0.2) -> str:
        """Call API with dummy tool to work around empty content bug."""
        import httpx as _httpx

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }

        payload = {
            "model": self.model,
            "messages": messages,
            "tools": [self.DUMMY_TOOL],
            "max_tokens": max_tokens,
            "temperature": temperature,
        }

        resp = _httpx.post(
            f"{self.base_url}/chat/completions",
            headers=headers,
            json=payload,
            timeout=120.0
        )
        data = resp.json()
        msg = data.get("choices", [{}])[0].get("message", {})

        # Case 1: Model responded with content directly
        content = msg.get("content") or ""
        if content.strip():
            return content

        # Case 2: Model called our dummy tool — extract the answer
        tool_calls = msg.get("tool_calls") or []
        if tool_calls:
            for tc in tool_calls:
                func = tc.get("function", {})
                if func.get("name") == "respond_to_user":
                    try:
                        args = json.loads(func.get("arguments", "{}"))
                        answer = args.get("answer", "")
                        if answer:
                            print(f"[YTL] Got answer via tool call ({len(answer)} chars)")
                            return answer
                    except json.JSONDecodeError:
                        # Arguments might be raw text
                        return func.get("arguments", "")

                # Other tool called — send tool result and get final answer
                tool_result_messages = messages + [
                    msg,  # assistant message with tool call
                    {
                        "role": "tool",
                        "tool_call_id": tc.get("id", ""),
                        "content": "Tool not available. Please answer the question directly based on the context provided."
                    }
                ]
                # Second call without tools
                resp2 = _httpx.post(
                    f"{self.base_url}/chat/completions",
                    headers=headers,
                    json={
                        "model": self.model,
                        "messages": tool_result_messages,
                        "max_tokens": max_tokens,
                        "temperature": temperature,
                    },
                    timeout=120.0
                )
                data2 = resp2.json()
                content2 = data2.get("choices", [{}])[0].get("message", {}).get("content", "")
                if content2:
                    return content2

        # Case 3: Nothing worked
        print(f"[YTL] Warning: empty response even with tool workaround")
        return content

    def generate(self, prompt: str) -> str:
        if self.use_tool_workaround and self.api_key:
            messages = [
                {"role": "system", "content": "You are a helpful AI assistant. Always provide detailed and helpful answers. If you use the respond_to_user tool, put your COMPLETE answer in the 'answer' parameter."},
                {"role": "user", "content": prompt}
            ]
            return self._call_with_tool_workaround(messages)

        # Standard path (Ollama, other providers)
        client = self._get_client()
        response = client.chat.completions.create(
            model=self.model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000,
            temperature=0.2
        )
        return response.choices[0].message.content or ""

    def generate_stream(self, prompt: str) -> Generator[str, None, None]:
        if self.use_tool_workaround and self.api_key:
            # Tool workaround doesn't support streaming — fall back to non-streaming
            result = self.generate(prompt)
            # Simulate streaming by yielding chunks
            chunk_size = 50
            for i in range(0, len(result), chunk_size):
                yield result[i:i + chunk_size]
            return

        # Standard streaming path
        client = self._get_client()
        stream = client.chat.completions.create(
            model=self.model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000,
            temperature=0.2,
            stream=True
        )
        for chunk in stream:
            if chunk.choices and chunk.choices[0].delta.content:
                yield chunk.choices[0].delta.content


# ============================================================================
# ULTRA MODE PROVIDERS
# ============================================================================

class ConversationMemory:
    """SQLite-based conversation memory for ultra mode."""

    def __init__(self, db_path: str = "/opt/jkst-master-ai/logs/memory.db"):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        conn = sqlite3.connect(self.db_path)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS conversations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id TEXT,
                role TEXT,
                content TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS query_feedback (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                query TEXT,
                expanded_queries TEXT,
                strategy TEXT,
                self_eval_score REAL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.commit()
        conn.close()

    def save_turn(self, session_id: str, role: str, content: str):
        conn = sqlite3.connect(self.db_path)
        conn.execute(
            "INSERT INTO conversations (session_id, role, content) VALUES (?, ?, ?)",
            (session_id, role, content)
        )
        conn.commit()
        conn.close()

    def get_history(self, session_id: str, limit: int = 10) -> List[Dict[str, str]]:
        conn = sqlite3.connect(self.db_path)
        rows = conn.execute(
            "SELECT role, content FROM conversations WHERE session_id = ? ORDER BY id DESC LIMIT ?",
            (session_id, limit * 2)
        ).fetchall()
        conn.close()
        return [{"role": r, "content": c} for r, c in reversed(rows)]


class UltraRetriever(BaseRetriever):
    """
    Ultra retriever with:
    1. Query expansion (LLM generates 3 search variants)
    2. Hybrid search (vector + BM25 keyword)
    3. Multi-source fusion
    """

    def __init__(self, generator: BaseGenerator):
        self.generator = generator
        self.top_k = int(os.environ.get("ULTRA_RETRIEVAL_TOP_K", "15"))

    def expand_query(self, query: str) -> List[str]:
        """Use LLM to generate query variants for better recall."""
        prompt = f"""Anda adalah pakar pencarian maklumat. Tulis 3 versi berbeza soalan berikut untuk carian dokumen.
Setiap versi perlu menggunakan kata kunci dan frasa yang berbeza tetapi bermaksud sama.
Fokus pada istilah Bahasa Melayu dan istilah undang-undang syariah.

Soalan asal: {query}

Jawab HANYA dengan 3 baris, satu soalan setiap baris. Tiada nombor atau bullet points:"""

        try:
            result = self.generator.generate(prompt)
            variants = [v.strip() for v in result.strip().split('\n') if v.strip()]
            # Always include original query
            all_queries = [query] + variants[:3]
            print(f"[ULTRA] Query expansion: {query} → {len(all_queries)} variants")
            for i, q in enumerate(all_queries):
                print(f"  [{i}] {q[:80]}")
            return all_queries
        except Exception as e:
            print(f"[ULTRA] Query expansion failed: {e}")
            return [query]

    async def retrieve(self, query: str) -> Tuple[List[str], List[Dict[str, Any]]]:
        """Hybrid retrieval: vector search + BM25 across expanded queries."""
        # Ensure knowledge is ingested
        await asyncio.to_thread(ingest_knowledge_to_chroma)

        # Step 1: Expand query
        expanded = await asyncio.to_thread(self.expand_query, query)

        collection = get_chroma_collection()
        model = get_embedding_model()

        # Step 2: Vector search across all query variants
        all_vector_results = {}  # doc_id -> (text, metadata, best_score)

        for q in expanded:
            q_embedding = model.encode(q).tolist()
            results = collection.query(
                query_embeddings=[q_embedding],
                n_results=self.top_k,
                include=["documents", "metadatas", "distances"]
            )

            if results and results['documents']:
                for doc, meta, dist in zip(
                    results['documents'][0],
                    results['metadatas'][0],
                    results['distances'][0]
                ):
                    doc_id = hashlib.md5(doc[:200].encode()).hexdigest()
                    score = 1.0 - dist
                    if doc_id not in all_vector_results or score > all_vector_results[doc_id][2]:
                        all_vector_results[doc_id] = (doc, meta, score)

        # Step 3: BM25 keyword search
        bm25_scores = {}
        try:
            from rank_bm25 import BM25Okapi

            # Get all documents from collection for BM25
            all_docs = collection.get(include=["documents", "metadatas"])
            if all_docs and all_docs['documents']:
                tokenized = [d.lower().split() for d in all_docs['documents']]
                bm25 = BM25Okapi(tokenized)

                # Score with original query
                query_tokens = query.lower().split()
                scores = bm25.get_scores(query_tokens)

                for i, (doc, meta, score) in enumerate(zip(
                    all_docs['documents'],
                    all_docs['metadatas'],
                    scores
                )):
                    if score > 0:
                        doc_id = hashlib.md5(doc[:200].encode()).hexdigest()
                        bm25_scores[doc_id] = (doc, meta, float(score))

            print(f"[ULTRA] BM25 found {len(bm25_scores)} keyword matches")
        except Exception as e:
            print(f"[ULTRA] BM25 search failed: {e}")

        # Step 4: Fuse results (Reciprocal Rank Fusion)
        fused = {}

        # Add vector results
        vector_sorted = sorted(all_vector_results.items(), key=lambda x: x[1][2], reverse=True)
        for rank, (doc_id, (doc, meta, score)) in enumerate(vector_sorted):
            rrf_score = 1.0 / (60 + rank)  # k=60 is standard RRF constant
            fused[doc_id] = {
                "doc": doc, "meta": meta,
                "vector_score": score, "bm25_score": 0,
                "rrf_score": rrf_score
            }

        # Add BM25 results
        bm25_sorted = sorted(bm25_scores.items(), key=lambda x: x[1][2], reverse=True)
        for rank, (doc_id, (doc, meta, score)) in enumerate(bm25_sorted):
            rrf_addition = 1.0 / (60 + rank)
            if doc_id in fused:
                fused[doc_id]["bm25_score"] = score
                fused[doc_id]["rrf_score"] += rrf_addition  # Boost docs found by both
            else:
                fused[doc_id] = {
                    "doc": doc, "meta": meta,
                    "vector_score": 0, "bm25_score": score,
                    "rrf_score": rrf_addition
                }

        # Sort by fused RRF score
        final_sorted = sorted(fused.items(), key=lambda x: x[1]["rrf_score"], reverse=True)

        contexts = []
        sources = []
        for doc_id, data in final_sorted[:self.top_k]:
            filename = data["meta"].get("filename", "Unknown")
            contexts.append(f"[SUMBER - DOKUMEN TEMPATAN: {filename}]\n{data['doc']}")
            sources.append({
                "type": "Dokumen (Hybrid Search)",
                "filename": filename,
                "page_content": data["doc"][:500],
                "score": round(data["rrf_score"], 4),
                "vector_score": round(data["vector_score"], 4),
                "bm25_score": round(data["bm25_score"], 4),
                "source_uri": f"local://{filename}",
                "priority": "PRIMARY"
            })

        print(f"[ULTRA] Hybrid retrieval: {len(contexts)} results "
              f"(vector: {len(all_vector_results)}, BM25: {len(bm25_scores)}, fused: {len(fused)})")

        return contexts, sources


class UltraReranker(BaseReranker):
    """Cross-encoder reranking using sentence-transformers."""

    def __init__(self):
        self._cross_encoder = None

    def _get_model(self):
        if self._cross_encoder is None:
            from sentence_transformers import CrossEncoder
            model_name = os.environ.get(
                "ULTRA_RERANKER_MODEL",
                "cross-encoder/ms-marco-MiniLM-L-6-v2"
            )
            print(f"[ULTRA] Loading cross-encoder: {model_name}")
            self._cross_encoder = CrossEncoder(model_name)
        return self._cross_encoder

    def rerank(self, query: str, documents: List[Dict[str, Any]], top_n: int = 5) -> List[Dict[str, Any]]:
        if not documents or len(documents) <= top_n:
            return documents

        try:
            model = self._get_model()

            # Prepare query-document pairs
            pairs = [
                (query, doc.get("page_content", "")[:512])
                for doc in documents
            ]

            scores = model.predict(pairs)

            # Sort by cross-encoder score
            scored = list(zip(scores, documents))
            scored.sort(key=lambda x: x[0], reverse=True)

            result = []
            for score, doc in scored[:top_n]:
                doc_copy = doc.copy()
                doc_copy["rerank_score"] = float(score)
                result.append(doc_copy)

            print(f"[ULTRA] Cross-encoder reranking: {len(documents)} → {len(result)} docs")
            for i, doc in enumerate(result[:3]):
                print(f"  #{i+1}: {doc.get('filename', '?')[:40]} (score: {doc.get('rerank_score', 0):.4f})")

            return result

        except Exception as e:
            print(f"[ULTRA] Cross-encoder reranking failed: {e}")
            return documents[:top_n]


class UltraEnhancer:
    """
    Ultra mode enhancements:
    - Self-evaluation
    - Follow-up suggestions
    - Answer quality scoring
    """

    def __init__(self, generator: BaseGenerator):
        self.generator = generator
        self.memory = ConversationMemory()

    def self_evaluate(self, query: str, answer: str, contexts: List[str]) -> Dict[str, Any]:
        """LLM evaluates its own answer quality."""
        context_preview = "\n".join(c[:200] for c in contexts[:3])

        eval_prompt = f"""Anda adalah pemeriksa kualiti jawapan AI. Nilai jawapan berikut:

SOALAN: {query}

KONTEKS YANG ADA:
{context_preview}

JAWAPAN AI:
{answer[:1000]}

Nilai 1-5 untuk setiap kriteria:
1. RELEVAN: Adakah jawapan menjawab soalan? (1=tidak relevan, 5=sangat relevan)
2. TEPAT: Adakah jawapan berdasarkan konteks/dokumen? (1=reka sendiri, 5=berdasarkan sumber)
3. LENGKAP: Adakah jawapan cukup lengkap? (1=tidak lengkap, 5=sangat lengkap)

Jawab HANYA dalam format JSON:
{{"relevan": N, "tepat": N, "lengkap": N, "purata": N.N, "nota": "satu ayat ulasan"}}"""

        try:
            result = self.generator.generate(eval_prompt)
            # Extract JSON from response
            json_match = re.search(r'\{[^}]+\}', result)
            if json_match:
                eval_data = json.loads(json_match.group())
                # Calculate average if not provided
                if "purata" not in eval_data:
                    scores = [eval_data.get("relevan", 3), eval_data.get("tepat", 3), eval_data.get("lengkap", 3)]
                    eval_data["purata"] = round(sum(scores) / len(scores), 1)
                print(f"[ULTRA] Self-eval: {eval_data.get('purata', '?')}/5 — {eval_data.get('nota', '')}")
                return eval_data
        except Exception as e:
            print(f"[ULTRA] Self-evaluation failed: {e}")

        return {"relevan": 3, "tepat": 3, "lengkap": 3, "purata": 3.0, "nota": "Penilaian automatik gagal"}

    def suggest_followups(self, query: str, answer: str) -> List[str]:
        """Generate follow-up question suggestions."""
        prompt = f"""Berdasarkan perbualan berikut, cadangkan 3 soalan susulan yang mungkin berguna.

Soalan pengguna: {query}
Jawapan ringkas: {answer[:500]}

Tulis 3 soalan susulan dalam Bahasa Melayu. Satu soalan setiap baris. Tiada nombor:"""

        try:
            result = self.generator.generate(prompt)
            suggestions = [s.strip() for s in result.strip().split('\n') if s.strip() and len(s.strip()) > 10]
            return suggestions[:3]
        except Exception as e:
            print(f"[ULTRA] Follow-up suggestion failed: {e}")
            return []

    def build_ultra_prompt(self, query: str, contexts: List[str], web_contexts: List[str],
                           history: List[Dict[str, str]], query_type: str = "hybrid",
                           local_contexts: List[str] = None,
                           jkst_news_contexts: List[str] = None,
                           expanded_queries: List[str] = None,
                           eval_data: Dict[str, Any] = None) -> str:
        """Build enhanced prompt for ultra mode with chain-of-thought."""

        # Standard context building (same as base)
        rag_text = "\n\n---\n\n".join(contexts) if contexts else ""
        web_text = "\n\n---\n\n".join(web_contexts) if web_contexts else ""
        local_text = "\n\n---\n\n".join(local_contexts) if local_contexts else ""

        history_text = ""
        if len(history) > 1:
            for msg in history[:-1]:
                role = "Pengguna" if msg["role"] == "user" else "Pembantu"
                history_text += f"{role}: {msg['content']}\n"

        # Ultra-enhanced system prompt
        system_prompt = """Anda adalah Pembantu AI JKST (Jabatan Kehakiman Syariah Terengganu) — versi ULTRA.

ANDA MESTI mengikut proses pemikiran berstruktur (Chain of Thought):

1. FAHAMI: Apakah sebenarnya yang ditanya? Kenal pasti niat sebenar pengguna.
2. ANALISIS: Semak semua konteks/dokumen yang ada. Kenal pasti maklumat yang relevan.
3. HUBUNGKAN: Gabungkan maklumat dari pelbagai sumber jika perlu.
4. JAWAB: Berikan jawapan yang lengkap, tersusun, dan mudah difahami.
5. SUMBER: Nyatakan sumber dengan jelas. Jika maklumat tidak pasti, nyatakan.

PERATURAN KEUTAMAAN SUMBER:
1. Dokumen Rasmi JKST → SUMBER UTAMA
2. Fail Pengetahuan Tempatan → SUMBER KEDUA
3. Maklumat Web → SUMBER TAMBAHAN sahaja

FORMAT JAWAPAN:
- Gunakan heading (##) untuk bahagian utama
- Gunakan senarai bernombor untuk langkah/prosedur
- Gunakan senarai bullet untuk maklumat umum
- Sertakan pautan muat turun jika ada
- Akhiri dengan "Soalan Berkaitan" jika sesuai

AMARAN: JANGAN berikan nasihat undang-undang/agama. Arahkan kepada pegawai yang berkenaan.

MAKLUMAT HUBUNGAN JKST:
- Alamat: Tingkat 5, Bangunan Mahkamah Syariah, Jalan Sultan Mohamad, 21100 Kuala Terengganu
- Waktu: Ahad-Rabu: 8:00 AM – 4:00 PM, Khamis: 8:00 AM – 3:00 PM
- Telefon: 09-623 2323 | Emel: jkstr@esyariah.gov.my
- Web: https://syariah.terengganu.gov.my"""

        # Build context section
        context_section = ""
        if rag_text:
            context_section += f"\n=== DOKUMEN UTAMA ===\n{rag_text}\n"
        if local_text:
            context_section += f"\n=== PENGETAHUAN TEMPATAN ===\n{local_text}\n"
        if web_text:
            context_section += f"\n=== MAKLUMAT WEB (TAMBAHAN) ===\n{web_text}\n"
        if not context_section:
            context_section = "Tiada dokumen berkaitan dijumpai."

        # Add expanded queries info for transparency
        expansion_note = ""
        if expanded_queries and len(expanded_queries) > 1:
            expansion_note = f"\n[Sistem telah mencari dengan {len(expanded_queries)} variasi soalan untuk hasil yang lebih menyeluruh]\n"

        prompt = f"""{system_prompt}

{expansion_note}
{"SEJARAH PERBUALAN:" + chr(10) + history_text if history_text else ""}

{context_section}

SOALAN: {query}

Gunakan proses pemikiran berstruktur (fahami → analisis → hubungkan → jawab → sumber).
Berikan jawapan yang LENGKAP dan TERSUSUN:"""

        return prompt


# ============================================================================
# MODE FACTORY
# ============================================================================
_providers = {}


def get_providers() -> Dict[str, Any]:
    """Get providers for the current RAG_MODE. Lazy-initialized."""
    global _providers
    mode = RAG_MODE

    if mode not in _providers:
        print(f"\n{'='*60}")
        print(f"  INITIALIZING RAG MODE: {mode.upper()}")
        print(f"{'='*60}\n")

        # Determine which LLM generator to use for local/ultra modes
        # LLM_PROVIDER: "gemini" (default) or "openai" (YTL Ilmu, Ollama, etc.)
        llm_provider = os.environ.get("LLM_PROVIDER", "gemini").lower()

        def _create_generator(for_mode: str) -> BaseGenerator:
            """Create the appropriate LLM generator based on LLM_PROVIDER env var."""
            if for_mode == "google":
                # Google mode always uses Gemini
                return GeminiGenerator()

            if llm_provider == "openai":
                base_url = os.environ.get("LOCAL_LLM_BASE_URL", "")
                api_key = os.environ.get("LOCAL_LLM_API_KEY", "")
                model = os.environ.get("LOCAL_LLM_MODEL", "")
                if base_url and api_key:
                    print(f"[{for_mode.upper()}] Using OpenAI-compatible LLM: {base_url} / {model}")
                    return OpenAICompatibleGenerator()
                else:
                    print(f"[{for_mode.upper()}] OpenAI LLM requested but no LOCAL_LLM_BASE_URL/API_KEY set, falling back to Gemini")
                    return GeminiGenerator()
            else:
                print(f"[{for_mode.upper()}] Using Gemini LLM")
                return GeminiGenerator()

        if mode == "google":
            _providers[mode] = {
                "mode": "google",
                "retriever": None,  # Uses app.py's retrieve_from_rag()
                "reranker": None,   # Uses app.py's rerank_documents()
                "generator": _create_generator("google"),
                "enhancer": None,
                "description": "Google Cloud (Vertex AI RAG + Gemini + Cohere Rerank)"
            }

        elif mode == "local":
            generator = _create_generator("local")
            llm_name = "YTL Ilmu" if llm_provider == "openai" else "Gemini"
            _providers[mode] = {
                "mode": "local",
                "retriever": LocalRetriever(),
                "reranker": LocalReranker(),
                "generator": generator,
                "enhancer": None,
                "description": f"Local (ChromaDB + Mesolitica Malay Embeddings + {llm_name} LLM)"
            }

        elif mode == "ultra":
            generator = _create_generator("ultra")
            llm_name = "YTL Ilmu" if llm_provider == "openai" else "Gemini"
            enhancer = UltraEnhancer(generator)
            _providers[mode] = {
                "mode": "ultra",
                "retriever": UltraRetriever(generator),
                "reranker": UltraReranker(),
                "generator": generator,
                "enhancer": enhancer,
                "memory": enhancer.memory,
                "description": f"Ultra (Hybrid Search + Query Expansion + Cross-Encoder Rerank + Self-Eval + {llm_name} LLM)"
            }

        else:
            raise ValueError(f"Unknown RAG_MODE: {mode}. Use 'google', 'local', or 'ultra'.")

        print(f"[{mode.upper()}] Providers initialized: {_providers[mode]['description']}")

    return _providers[mode]


def get_mode_info() -> Dict[str, Any]:
    """Get current mode information for the /api/mode endpoint."""
    providers = get_providers()
    mode = RAG_MODE

    # Determine active LLM name
    llm_provider = os.environ.get("LLM_PROVIDER", "gemini").lower()
    if llm_provider == "openai" and mode != "google":
        llm_model = os.environ.get("LOCAL_LLM_MODEL", "unknown")
        llm_base = os.environ.get("LOCAL_LLM_BASE_URL", "")
        # Extract provider name from URL
        if "ytlailabs" in llm_base:
            llm_name = f"YTL Ilmu ({llm_model})"
        elif "localhost" in llm_base or "127.0.0.1" in llm_base:
            llm_name = f"Ollama ({llm_model})"
        else:
            llm_name = f"OpenAI-compatible ({llm_model})"
    else:
        llm_name = "Gemini 2.0 Flash"

    info = {
        "current_mode": mode,
        "description": providers["description"],
        "llm_provider": llm_provider if mode != "google" else "gemini",
        "features": {}
    }

    if mode == "google":
        info["features"] = {
            "retrieval": "Google Vertex AI RAG Corpus",
            "embeddings": "Google (via RAG API)",
            "reranking": "Cohere rerank-v3.5 / Gemini",
            "llm": "Gemini 2.0 Flash",
            "storage": "Google Cloud Storage",
            "query_expansion": False,
            "hybrid_search": False,
            "self_evaluation": False,
            "conversation_memory": False,
            "followup_suggestions": False,
            "chain_of_thought": False,
        }
    elif mode == "local":
        info["features"] = {
            "retrieval": "ChromaDB (local vector DB)",
            "embeddings": "mesolitica/mistral-embedding-191m-8k (Malay)",
            "reranking": "Keyword overlap (local)",
            "llm": llm_name,
            "storage": "Local filesystem",
            "query_expansion": False,
            "hybrid_search": False,
            "self_evaluation": False,
            "conversation_memory": False,
            "followup_suggestions": False,
            "chain_of_thought": False,
        }
    elif mode == "ultra":
        info["features"] = {
            "retrieval": "ChromaDB + BM25 (hybrid)",
            "embeddings": "mesolitica/mistral-embedding-191m-8k (Malay)",
            "reranking": "Cross-encoder (ms-marco-MiniLM-L-6-v2)",
            "llm": llm_name,
            "storage": "Local filesystem",
            "query_expansion": True,
            "hybrid_search": True,
            "self_evaluation": True,
            "conversation_memory": True,
            "followup_suggestions": True,
            "chain_of_thought": True,
        }

    return info
